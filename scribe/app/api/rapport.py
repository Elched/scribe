"""
api/rapport.py — Génération rapport de clôture DOCX + REX SCRIBE v5
"""
import json, io
from datetime import datetime, timezone
from typing import Optional, List
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.database import get_db
from app.api.auth import get_current_user
from app.models import SitrepEntry, Decision, Presence, Consigne, RexEntry, Task

router = APIRouter()

URG_LABELS = {1: "VEILLE", 2: "ALERTE", 3: "CRISE", 4: "CRITIQUE"}
TYPE_COLORS = {"CYBER": RGBColor(0x3b, 0x82, 0xf6), "SANITAIRE": RGBColor(0x10, 0xb9, 0x81), "MIXTE": RGBColor(0xf5, 0x9e, 0x0b)}
URG_COLORS  = {1: RGBColor(0x4a, 0xde, 0x80), 2: RGBColor(0xfb, 0xbf, 0x24), 3: RGBColor(0xef, 0x44, 0x44), 4: RGBColor(0xc0, 0x84, 0xfc)}

def _heading(doc, text, level=1, color=RGBColor(0x1e, 0x40, 0xaf)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
    return p

def _add_table_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, (cell, val) in enumerate(zip(row.cells, cells)):
        cell.text = str(val)
        if bold_first and i == 0:
            cell.paragraphs[0].runs[0].bold = True

def _minutes_to_str(m):
    if m is None: return "N/A"
    h, mn = divmod(m, 60)
    return f"{h}h{mn:02d}" if h else f"{mn}min"


def generate_rapport_docx(incident_id: int, db: Session) -> bytes:
    inc = db.query(SitrepEntry).filter(SitrepEntry.id == incident_id).first()
    if not inc:
        raise HTTPException(404, "Incident non trouvé")

    decisions = db.query(Decision).order_by(Decision.timestamp).all()
    presences = db.query(Presence).order_by(Presence.timestamp).all()
    consignes = db.query(Consigne).order_by(Consigne.timestamp).all()

    # Calculs métriques
    now = datetime.now(timezone.utc)
    ts  = inc.timestamp.replace(tzinfo=timezone.utc) if inc.timestamp.tzinfo is None else inc.timestamp
    duree_min = int((inc.resolved_at.replace(tzinfo=timezone.utc) - ts).total_seconds() / 60) if inc.resolved_at else None
    jalons_data = json.loads(inc.jalons) if inc.jalons else []
    nb_done = sum(1 for j in jalons_data if j.get("done"))

    doc = Document()

    # ── Page de garde ─────────────────────────────────────
    section = doc.sections[0]
    section.page_height = Cm(29.7); section.page_width = Cm(21)
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SCRIBE v5 — Établissement")
    run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RAPPORT DE CLÔTURE D'INCIDENT")
    run.font.size = Pt(22); run.bold = True
    run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(inc.fait[:120])
    run.font.size = Pt(14); run.italic = True

    doc.add_paragraph()
    # Tableau récap
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows_data = [
        ("Référence incident", f"#{inc.id}"),
        ("Date/heure déclaration", ts.strftime("%d/%m/%Y %H:%M")),
        ("Type de crise", inc.type_crise),
        ("Niveau d'urgence", URG_LABELS.get(inc.urgency, str(inc.urgency))),
        ("Site", inc.site_id),
        ("UF concernée", inc.unite_fonctionnelle or "—"),
        ("Directeur de crise", inc.directeur_crise or "—"),
        ("Déclarant", inc.declarant_nom),
        ("Statut final", inc.status),
        ("Date/heure résolution", inc.resolved_at.strftime("%d/%m/%Y %H:%M") if inc.resolved_at else "Non résolu"),
        ("Durée totale", _minutes_to_str(duree_min)),
        ("Jalons complétés", f"{nb_done}/{len(jalons_data)}"),
    ]
    for k, v in rows_data:
        row = tbl.add_row()
        row.cells[0].text = k; row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = v

    doc.add_page_break()

    # ── 1. Description de l'incident ─────────────────────
    _heading(doc, "1. Description de l'incident", 1)
    _heading(doc, "1.1 Fait déclaré", 2)
    doc.add_paragraph(inc.fait)
    if inc.analyse:
        _heading(doc, "1.2 Analyse d'impact", 2)
        doc.add_paragraph(inc.analyse)
    if inc.moyens_engages:
        _heading(doc, "1.3 Moyens engagés", 2)
        doc.add_paragraph(inc.moyens_engages)
    if inc.actions_remediation:
        _heading(doc, "1.4 Actions de remédiation", 2)
        doc.add_paragraph(inc.actions_remediation)
    if inc.intervenant_nom:
        _heading(doc, "1.5 Intervenants", 2)
        doc.add_paragraph(f"{inc.intervenant_nom}" + (f" — {inc.intervenant_contact}" if inc.intervenant_contact else ""))

    # ── 2. Chronologie des jalons ─────────────────────────
    if jalons_data:
        _heading(doc, "2. Chronologie des jalons de résolution", 1)
        tbl2 = doc.add_table(rows=1, cols=3)
        tbl2.style = "Table Grid"
        hdr = tbl2.rows[0].cells
        hdr[0].text = "Jalon"; hdr[1].text = "Statut"; hdr[2].text = "Complété le"
        for j in jalons_data:
            r = tbl2.add_row()
            r.cells[0].text = j.get("label", "")
            r.cells[1].text = "✓ FAIT" if j.get("done") else "— En attente"
            r.cells[2].text = j.get("done_at", "")[:16] if j.get("done_at") else "—"

    # ── 3. Décisions de cellule ───────────────────────────
    if decisions:
        _heading(doc, "3. Décisions de la cellule de crise", 1)
        tbl3 = doc.add_table(rows=1, cols=4)
        tbl3.style = "Table Grid"
        h = tbl3.rows[0].cells
        h[0].text="Horodatage"; h[1].text="Responsable"; h[2].text="Décision"; h[3].text="Base régl."
        for d in decisions:
            r = tbl3.add_row()
            r.cells[0].text = d.timestamp.strftime("%d/%m %H:%M") if d.timestamp else "—"
            r.cells[1].text = d.responsable or "—"
            r.cells[2].text = d.contenu
            r.cells[3].text = d.base_reglementaire or "—"

    # ── 4. Présences cellule ──────────────────────────────
    if presences:
        _heading(doc, "4. Registre des présences", 1)
        tbl4 = doc.add_table(rows=1, cols=4)
        tbl4.style = "Table Grid"
        h = tbl4.rows[0].cells
        h[0].text="Horodatage"; h[1].text="Nom"; h[2].text="Rôle"; h[3].text="Action"
        for p in presences:
            r = tbl4.add_row()
            r.cells[0].text = p.timestamp.strftime("%d/%m %H:%M") if p.timestamp else "—"
            r.cells[1].text = p.nom
            r.cells[2].text = p.role or "—"
            r.cells[3].text = p.action

    # ── 5. Consignes de relève ────────────────────────────
    if consignes:
        _heading(doc, "5. Consignes de relève", 1)
        for c in consignes:
            p_c = doc.add_paragraph()
            p_c.add_run(f"[{c.timestamp.strftime('%d/%m %H:%M') if c.timestamp else '—'}] Pour {c.pour} : ").bold = True
            p_c.add_run(c.texte)
            if c.accuse:
                p_c.add_run(f" ✓ Accusé le {c.accuse_at.strftime('%d/%m %H:%M') if c.accuse_at else '?'}").italic = True

    # ── 6. Avis Albert ───────────────────────────────────
    if inc.albert_avis:
        _heading(doc, "6. Avis Albert AI", 1)
        doc.add_paragraph(inc.albert_avis)
        doc.add_paragraph("Note : analyse générée automatiquement par Albert AI (DINUM/Etalab)."
                          "\nDoit être interprétée par un expert métier avant toute décision.", style="Caption")

    # ── 7. Métriques ─────────────────────────────────────
    _heading(doc, "7. Métriques de gestion", 1)
    tbl_m = doc.add_table(rows=0, cols=2)
    tbl_m.style = "Table Grid"
    metrics = [
        ("Durée totale (déclaration → résolution)", _minutes_to_str(duree_min)),
        ("Jalons complétés", f"{nb_done}/{len(jalons_data)} ({int(nb_done/len(jalons_data)*100) if jalons_data else 0}%)"),
        ("Nombre de décisions actées", str(len(decisions))),
        ("Intervenants cellule", str(len(set(p.nom for p in presences)))),
        ("Consignes de relève", str(len(consignes))),
    ]
    for k, v in metrics:
        r = tbl_m.add_row()
        r.cells[0].text = k; r.cells[0].paragraphs[0].runs[0].bold = True
        r.cells[1].text = v

    # ── Pied de page ─────────────────────────────────────
    doc.add_paragraph()
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_foot.add_run(f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} — SCRIBE v5 Établissement — CONFIDENTIEL")
    run.font.size = Pt(8); run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── REX ──────────────────────────────────────────────────

class RexCreate(BaseModel):
    incident_id:     Optional[int] = None
    titre:           str
    type_crise:      Optional[str] = None
    duree_minutes:   Optional[int] = None
    nb_poles:        int = 0
    nb_decisions:    int = 0
    nb_jalons_total: int = 0
    nb_jalons_done:  int = 0
    mttd_minutes:    Optional[int] = None
    mttr_minutes:    Optional[int] = None
    points_positifs: Optional[List[str]] = []
    points_amelio:   Optional[List[str]] = []
    actions_futures: Optional[List[str]] = []
    lecons:          Optional[str] = ""
    redacteur:       Optional[str] = ""

class RexOut(BaseModel):
    id: int; incident_id: Optional[int]; titre: str; type_crise: Optional[str]
    created_at: Optional[datetime]; duree_minutes: Optional[int]
    nb_poles: int; nb_decisions: int; nb_jalons_total: int; nb_jalons_done: int
    mttd_minutes: Optional[int]; mttr_minutes: Optional[int]
    points_positifs: Optional[str]; points_amelio: Optional[str]
    actions_futures: Optional[str]; lecons: Optional[str]; redacteur: Optional[str]
    class Config: from_attributes = True


@router.get("/rapport/{incident_id}")
def download_rapport(incident_id: int, db: Session = Depends(get_db)):
    docx_bytes = generate_rapport_docx(incident_id, db)
    filename = f"rapport_cloture_incident_{incident_id}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@router.get("/rex", response_model=List[RexOut])
def list_rex(db: Session = Depends(get_db)):
    return db.query(RexEntry).order_by(RexEntry.created_at.desc()).all()

@router.post("/rex", response_model=RexOut)
def create_rex(body: RexCreate, db: Session = Depends(get_db)):
    data = body.dict()
    for k in ["points_positifs", "points_amelio", "actions_futures"]:
        if isinstance(data[k], list):
            data[k] = json.dumps(data[k], ensure_ascii=False)
    r = RexEntry(**data)
    db.add(r); db.commit(); db.refresh(r)
    return r

@router.delete("/rex/{rid}")
def delete_rex(rid: int, db: Session = Depends(get_db)):
    r = db.query(RexEntry).filter(RexEntry.id == rid).first()
    if not r: raise HTTPException(404, "REX non trouvé")
    db.delete(r); db.commit()
    return {"status": "deleted"}

@router.get("/rex-stats")
def rex_stats(db: Session = Depends(get_db)):
    entries = db.query(RexEntry).all()
    if not entries:
        return {"total": 0}
    mttrs = [e.mttr_minutes for e in entries if e.mttr_minutes]
    mttds = [e.mttd_minutes for e in entries if e.mttd_minutes]
    by_type = {}
    for e in entries:
        t = e.type_crise or "INCONNU"
        by_type[t] = by_type.get(t, 0) + 1
    return {
        "total": len(entries),
        "avg_mttr_min": int(sum(mttrs)/len(mttrs)) if mttrs else None,
        "avg_mttd_min": int(sum(mttds)/len(mttds)) if mttds else None,
        "by_type": by_type,
        "avg_jalons_pct": int(sum(e.nb_jalons_done/(e.nb_jalons_total or 1)*100 for e in entries)/len(entries)),
    }

# ── Export main courante complet ─────────────────────────────────────────

@router.get("/export-main-courante")
def export_main_courante(db: Session = Depends(get_db)):
    """Export CSV chronologique complet : incidents, décisions, présences,
    relève, kanban, communiqués, REX."""
    import csv, io as _io
    from app.api.status_page import StatusPageChronologie
    from app.models import RexEntry

    output = _io.StringIO()
    w = csv.writer(output, delimiter=";", quotechar='"', quoting=csv.QUOTE_ALL)
    w.writerow(["Horodatage", "Catégorie", "Sous-type", "Acteur", "Contenu", "Détail"])

    def fmt(dt):
        if dt is None: return ""
        try: return dt.strftime("%d/%m/%Y %H:%M")
        except Exception: return str(dt)

    events = []  # (datetime_obj, row_tuple)

    # ── Incidents ────────────────────────────────────────────────────────
    for i in db.query(SitrepEntry).order_by(SitrepEntry.timestamp).all():
        events.append((i.timestamp, [
            fmt(i.timestamp), "INCIDENT", i.type_crise or "",
            i.declarant_nom or "",
            i.fait or "",
            f"U{i.urgency or ''} | {i.status or ''} | {i.site_id or ''} | dir:{i.directeur_crise or ''}"
        ]))
        # Changement de statut via jalons
        try:
            for j in json.loads(i.jalons or "[]"):
                if j.get("done") and j.get("done_at"):
                    events.append((j["done_at"], [
                        j["done_at"][:16].replace("T"," "), "INCIDENT — JALON",
                        i.type_crise or "", f"Incident #{i.id}",
                        j.get("label",""), ""
                    ]))
        except Exception:
            pass

    # ── Décisions cellule ────────────────────────────────────────────────
    for d in db.query(Decision).order_by(Decision.timestamp).all():
        events.append((d.timestamp, [
            fmt(d.timestamp), "DÉCISION CELLULE", d.base_reglementaire or "",
            d.responsable or "", d.contenu or "",
            d.statut_validation or ""
        ]))

    # ── Entrées/sorties cellule ──────────────────────────────────────────
    for p in db.query(Presence).order_by(Presence.timestamp).all():
        events.append((p.timestamp, [
            fmt(p.timestamp), "PRÉSENCE CELLULE", p.action or "",
            p.nom or "", p.role or "", ""
        ]))

    # ── Relève / consignes ───────────────────────────────────────────────
    for c in db.query(Consigne).order_by(Consigne.timestamp).all():
        ack = f"Accusé par {c.accuse_par or '?'} à {fmt(c.accuse_at)}" if c.accuse else "Non accusé"
        events.append((c.timestamp, [
            fmt(c.timestamp), "RELÈVE — CONSIGNE", "TRANSMISE",
            c.pour or "", c.texte or "", ack
        ]))

    # ── Kanban ───────────────────────────────────────────────────────────
    for t in db.query(Task).order_by(Task.created_at).all():
        prio = {1:"BASSE",2:"NORMALE",3:"HAUTE",4:"CRITIQUE"}.get(t.priorite,"?")
        events.append((t.created_at, [
            fmt(t.created_at), "KANBAN", t.colonne or "",
            t.assignee or "", t.titre or "",
            f"{prio} | {t.description or ''}"
        ]))
        if t.updated_at and t.updated_at != t.created_at:
            events.append((t.updated_at, [
                fmt(t.updated_at), "KANBAN — DÉPLACEMENT", t.colonne or "",
                t.assignee or "", t.titre or "", ""
            ]))

    # ── REX ──────────────────────────────────────────────────────────────
    for r in db.query(RexEntry).order_by(RexEntry.created_at).all():
        events.append((r.created_at, [
            fmt(r.created_at), "REX", r.type_crise or "",
            r.redacteur or "", r.titre or "",
            f"MTTD:{r.mttd_minutes or ''}min MTTR:{r.mttr_minutes or ''}min"
        ]))

    # ── Déclarations capacitaires ────────────────────────────────────────
    try:
        refs_cap = {r.id: r for r in db.query(CapaciteReferentiel).all()}
        for d in db.query(CapaciteDeclaration).order_by(CapaciteDeclaration.horodatage).all():
            ref_cap = refs_cap.get(d.referentiel_id)
            service = ref_cap.service_nom if ref_cap else f"Service #{d.referentiel_id}"
            alertes = []
            if d.alerte_lits:     alertes.append("Alerte lits")
            if d.alerte_rh:       alertes.append("Alerte RH")
            if d.alerte_materiel: alertes.append("Alerte matériel")
            detail = (
                f"Lits: {d.statut_lits} (H:{d.lits_vides_h} F:{d.lits_vides_f} I:{d.lits_vides_i}) | "
                f"RH:{d.statut_rh} | Mat:{d.statut_materiel}"
                + (f" | {'/ '.join(alertes)}" if alertes else "")
            )
            events.append((d.horodatage, [
                fmt(d.horodatage), "CAPACITÉ — DÉCLARATION",
                d.point or "",
                d.redacteur or "",
                service,
                detail + (f" | {d.commentaire_general}" if d.commentaire_general else "")
            ]))
    except Exception:
        pass

    # ── Communiqués / chronologie publique ──────────────────────────────
    try:
        for c in db.query(StatusPageChronologie).order_by(StatusPageChronologie.timestamp).all():
            events.append((c.timestamp, [
                fmt(c.timestamp), "COMMUNIQUÉ PUBLIC", "CHRONOLOGIE",
                c.publie_par or "", c.texte or "", ""
            ]))
    except Exception:
        pass

    # ── Trier et écrire ──────────────────────────────────────────────────
    def sort_key(ev):
        dt = ev[0]
        if dt is None: return ""
        if hasattr(dt, "isoformat"): return dt.isoformat()
        return str(dt)

    events.sort(key=sort_key)
    for _, row in events:
        w.writerow(row)

    output.seek(0)
    now_str = datetime.now().strftime("%Y%m%d_%H%M")
    return StreamingResponse(
        iter([output.getvalue().encode("utf-8-sig")]),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename=main_courante_{now_str}.csv"}
    )

# ── Nouvelle crise : archive ZIP + reset ─────────────────────────────────

@router.post("/archiver-crise")
def archiver_crise(db: Session = Depends(get_db), user=Depends(get_current_user)):
    """Archive la crise courante en ZIP sans toucher au tableau de bord."""
    """Archive la crise courante en ZIP et remet à zéro le tableau de bord."""
    import zipfile, io as _io, csv
    from app.api.status_page import StatusPage as SPModel, StatusPageChronologie
    from app.models import Task, RexEntry, CapaciteDeclaration, CapaciteReferentiel

    now_str = datetime.now().strftime("%Y%m%d_%H%M%S")
    zip_path = f"archives/crise_{now_str}.zip"

    import os
    os.makedirs("archives", exist_ok=True)

    buf = _io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:

        def to_csv(rows, headers):
            out = _io.StringIO()
            w = csv.writer(out, delimiter=';', quotechar='"', quoting=csv.QUOTE_ALL)
            w.writerow(headers)
            for r in rows: w.writerow(r)
            return out.getvalue().encode('utf-8-sig')

        # Incidents
        incs = db.query(SitrepEntry).all()
        zf.writestr("incidents.csv", to_csv(
            [(i.id, i.timestamp, i.type_crise, i.urgency, i.fait, i.analyse,
              i.status, i.site_id, i.declarant_nom, i.directeur_crise) for i in incs],
            ["id","timestamp","type","urgency","fait","analyse","status","site","declarant","directeur"]
        ))

        # Décisions
        decs = db.query(Decision).all()
        zf.writestr("decisions.csv", to_csv(
            [(d.id, d.timestamp, d.responsable, d.contenu, d.base_reglementaire, d.statut_validation) for d in decs],
            ["id","timestamp","responsable","contenu","base_reglementaire","statut_validation"]
        ))

        # Présences
        pres = db.query(Presence).all()
        zf.writestr("presences.csv", to_csv(
            [(p.id, p.timestamp, p.nom, p.role, p.action) for p in pres],
            ["id","timestamp","nom","role","action"]
        ))

        # Relève
        cons = db.query(Consigne).all()
        zf.writestr("releve.csv", to_csv(
            [(c.id, c.timestamp, c.pour, c.texte, c.accuse, c.accuse_at, c.accuse_par) for c in cons],
            ["id","timestamp","pour","texte","accuse","accuse_at","accuse_par"]
        ))

        # Kanban
        tasks = db.query(Task).all()
        zf.writestr("kanban.csv", to_csv(
            [(t.id, t.created_at, t.titre, t.colonne, t.priorite, t.assignee, t.description) for t in tasks],
            ["id","created_at","titre","colonne","priorite","assignee","description"]
        ))

        # REX
        rex = db.query(RexEntry).all()
        zf.writestr("rex.csv", to_csv(
            [(r.id, r.created_at, r.titre, r.type_crise, r.redacteur) for r in rex],
            ["id","created_at","titre","type_crise","redacteur"]
        ))

        # Statuts publics / communiqués
        try:
            sps = db.query(SPModel).all()
            zf.writestr("communiques.csv", to_csv(
                [(s.id, s.site_id, s.site_nom, s.niveau_global, s.message_public, s.updated_at) for s in sps],
                ["id","site_id","site_nom","niveau","message","updated_at"]
            ))
            chrons = db.query(StatusPageChronologie).all()
            zf.writestr("chronologie_publique.csv", to_csv(
                [(c.id, c.timestamp, c.texte, c.publie_par) for c in chrons],
                ["id","timestamp","texte","publie_par"]
            ))
        except Exception:
            pass

        # Métadonnées de la crise
        meta = (
            f"SCRIBE — Archive de crise\n"
            f"Date d'archivage : {datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
            f"Archivé par : {getattr(user,'display_name','') or getattr(user,'username','')}\n"
            f"Incidents : {len(incs)}\n"
            f"Décisions : {len(decs)}\n"
            f"Tâches Kanban : {len(tasks)}\n"
        )
        zf.writestr("README.txt", meta)

    # Sauvegarder l'archive
    with open(zip_path, 'wb') as f:
        f.write(buf.getvalue())

    db.commit()

    return {
        "ok": True,
        "archive": zip_path,
        "message": f"Crise archivée dans {zip_path}."
    }


# ── Reset tableau de bord (séparé de l'archivage) ────────────────────────

@router.post("/reset-tableau-de-bord")
def reset_tableau_de_bord(db: Session = Depends(get_db), user=Depends(get_current_user)):
    """Remet le tableau de bord à zéro sans archiver."""
    from app.models import Task, RexEntry, CapaciteDeclaration, CapaciteReferentiel
    db.query(SitrepEntry).delete()
    db.query(Decision).delete()
    db.query(Presence).delete()
    db.query(Consigne).delete()
    db.query(Task).delete()
    try:
        from app.api.status_page import StatusPage as SPModel2, StatusPageChronologie as SPC2
        db.query(SPC2).delete()
        for sp in db.query(SPModel2).all():
            sp.published = False
            sp.message_public = ""
            sp.niveau_global = "OPERATIONNEL"
    except Exception:
        pass
    db.commit()
    return {"ok": True, "message": "Tableau de bord remis à zéro."}


# ── Garde la route nouvelle-crise pour compatibilité (archive + reset) ───

@router.post("/nouvelle-crise")
def nouvelle_crise(db: Session = Depends(get_db), user=Depends(get_current_user)):
    """Archive puis remet à zéro — appelle les deux routes en séquence."""
    archiver_crise(db=db, user=user)
    db.expunge_all()  # Rafraîchir la session
    from sqlalchemy.orm import Session as S
    reset_tableau_de_bord(db=db, user=user)
    return {"ok": True, "message": "Crise archivée et tableau de bord réinitialisé."}

# ── Rapport DOCX debriefing ───────────────────────────────────────────────

class DebriefEvent(BaseModel):
    ts: str = ""
    cat: str = ""
    acteur: str = ""
    contenu: str = ""
    note: str = ""

class DebriefRequest(BaseModel):
    metrics: dict = {}
    events: list = []
    annotations: dict = {}

@router.post("/debrief-docx")
def generate_debrief_docx(req: DebriefRequest, db: Session = Depends(get_db)):
    """Génère un rapport DOCX de debriefing depuis les données d'analyse."""
    try:
        import json as _j, os as _o
        _cfgp = _o.path.join(_o.path.dirname(__file__), "..", "static", "config.js")
        _cfgt = open(_cfgp, encoding="utf-8").read() if _o.path.exists(_cfgp) else ""
        _start = _cfgt.find("const SCRIBE_CONFIG = ") + len("const SCRIBE_CONFIG = ")
        _cfg = _j.loads(_cfgt[_start:_cfgt.rfind(";")]) if _start > 20 else {}
        nom_etab = _cfg.get("etablissement", {}).get("nom", "Établissement de santé")
    except Exception:
        nom_etab = "Établissement de santé"

    doc = Document()
    doc.core_properties.title = "Rapport de debriefing"

    # ── Titre ──
    p = doc.add_heading("RAPPORT DE DEBRIEFING DE CRISE", 0)
    p.runs[0].font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

    doc.add_heading(nom_etab, 1)
    doc.add_paragraph(f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
    doc.add_paragraph()

    # ── Métriques ──
    doc.add_heading("1. MÉTRIQUES CLÉS", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Indicateur"
    hdr[1].text = "Valeur"
    for k, v in req.metrics.items():
        row = table.add_row().cells
        labels = {
            "incidents": "Nombre d'incidents",
            "decisions": "Décisions prises",
            "presences": "Participants cellule",
            "kanban": "Tâches kanban",
        }
        row[0].text = labels.get(k, k)
        row[1].text = str(v)
    doc.add_paragraph()

    # ── Chronologie ──
    doc.add_heading("2. CHRONOLOGIE DES ÉVÉNEMENTS", 1)
    cat_colors = {
        "INCIDENT": RGBColor(0xef, 0x44, 0x44),
        "DÉCISION": RGBColor(0xf5, 0x9e, 0x0b),
        "PRÉSENCE": RGBColor(0x3d, 0x9e, 0xff),
        "KANBAN":   RGBColor(0xa8, 0x55, 0xf7),
        "RELÈVE":   RGBColor(0x6b, 0x74, 0x94),
        "COMMUNIQUÉ": RGBColor(0x00, 0xe5, 0xa0),
        "REX":      RGBColor(0x22, 0xd3, 0xee),
    }

    for ev in req.events[:60]:
        p = doc.add_paragraph()
        run_ts = p.add_run(f"[{ev.get('ts','')}] ")
        run_ts.font.size = Pt(8)
        run_ts.font.color.rgb = RGBColor(0x6b, 0x74, 0x94)

        cat = ev.get("cat","")
        color = next((v for k,v in cat_colors.items() if cat.startswith(k)), RGBColor(0x6b,0x74,0x94))
        run_cat = p.add_run(f"{cat} ")
        run_cat.font.size = Pt(8)
        run_cat.font.bold = True
        run_cat.font.color.rgb = color

        run_txt = p.add_run(ev.get("contenu",""))
        run_txt.font.size = Pt(9)

        if ev.get("acteur"):
            run_act = p.add_run(f"  — {ev['acteur']}")
            run_act.font.size = Pt(8)
            run_act.font.color.rgb = RGBColor(0x6b,0x74,0x94)

        # Annotation
        if ev.get("note"):
            p2 = doc.add_paragraph()
            r = p2.add_run(f"    📝 {ev['note']}")
            r.font.size = Pt(8)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0xf5, 0x9e, 0x0b)

    doc.add_paragraph()

    # ── Annotations ──
    all_notes = [(k,v) for k,v in req.annotations.items() if v.strip()]
    if all_notes:
        doc.add_heading("3. ANNOTATIONS ET POINTS CLÉS", 1)
        for ev_id, note in all_notes[:20]:
            p = doc.add_paragraph()
            p.add_run("• ").bold = True
            p.add_run(note)

    # ── Recommandations ──
    doc.add_heading("4. PLAN D'AMÉLIORATION", 1)
    doc.add_paragraph("À compléter par l'équipe lors du debriefing :")
    for label in ["Points forts à capitaliser :", "Axes d'amélioration identifiés :", "Actions à mettre en place avant la prochaine crise :"]:
        p = doc.add_paragraph(label)
        p.runs[0].bold = True
        for _ in range(3):
            doc.add_paragraph("___________________________________________")

    # ── Sauvegarde ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        iter([buf.getvalue()]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=debrief_crise.docx"}
    )
