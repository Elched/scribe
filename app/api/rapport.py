"""
api/rapport.py — Génération rapport de clôture DOCX + REX SCRIBE v5
"""
import json, io
from datetime import datetime, timezone
from typing import Optional, List
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.database import get_db
from app.models import SitrepEntry, Decision, Presence, Consigne, RexEntry

router = APIRouter()

URG_LABELS = {1: "VEILLE", 2: "ALERTE", 3: "CRISE", 4: "CRITIQUE"}
TYPE_COLORS = {"CYBER": RGBColor(0x3b, 0x82, 0xf6), "SANITAIRE": RGBColor(0x10, 0xb9, 0x81), "MIXTE": RGBColor(0xf5, 0x9e, 0x0b)}
URG_COLORS  = {1: RGBColor(0x4a, 0xde, 0x80), 2: RGBColor(0xfb, 0xbf, 0x24), 3: RGBColor(0xef, 0x44, 0x44), 4: RGBColor(0xc0, 0x84, 0xfc)}

def _heading(doc, text, level=1, color=RGBColor(0x1e, 0x40, 0xaf)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
    return p

def _add_table_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, (cell, val) in enumerate(zip(row.cells, cells)):
        cell.text = str(val)
        if bold_first and i == 0:
            cell.paragraphs[0].runs[0].bold = True

def _minutes_to_str(m):
    if m is None: return "N/A"
    h, mn = divmod(m, 60)
    return f"{h}h{mn:02d}" if h else f"{mn}min"


def generate_rapport_docx(incident_id: int, db: Session) -> bytes:
    inc = db.query(SitrepEntry).filter(SitrepEntry.id == incident_id).first()
    if not inc:
        raise HTTPException(404, "Incident non trouvé")

    decisions = db.query(Decision).order_by(Decision.timestamp).all()
    presences = db.query(Presence).order_by(Presence.timestamp).all()
    consignes = db.query(Consigne).order_by(Consigne.timestamp).all()

    # Calculs métriques
    now = datetime.now(timezone.utc)
    ts  = inc.timestamp.replace(tzinfo=timezone.utc) if inc.timestamp.tzinfo is None else inc.timestamp
    duree_min = int((inc.resolved_at.replace(tzinfo=timezone.utc) - ts).total_seconds() / 60) if inc.resolved_at else None
    jalons_data = json.loads(inc.jalons) if inc.jalons else []
    nb_done = sum(1 for j in jalons_data if j.get("done"))

    doc = Document()

    # ── Page de garde ─────────────────────────────────────
    section = doc.sections[0]
    section.page_height = Cm(29.7); section.page_width = Cm(21)
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SCRIBE v5 — Établissement")
    run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RAPPORT DE CLÔTURE D'INCIDENT")
    run.font.size = Pt(22); run.bold = True
    run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(inc.fait[:120])
    run.font.size = Pt(14); run.italic = True

    doc.add_paragraph()
    # Tableau récap
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows_data = [
        ("Référence incident", f"#{inc.id}"),
        ("Date/heure déclaration", ts.strftime("%d/%m/%Y %H:%M")),
        ("Type de crise", inc.type_crise),
        ("Niveau d'urgence", URG_LABELS.get(inc.urgency, str(inc.urgency))),
        ("Site", inc.site_id),
        ("UF concernée", inc.unite_fonctionnelle or "—"),
        ("Directeur de crise", inc.directeur_crise or "—"),
        ("Déclarant", inc.declarant_nom),
        ("Statut final", inc.status),
        ("Date/heure résolution", inc.resolved_at.strftime("%d/%m/%Y %H:%M") if inc.resolved_at else "Non résolu"),
        ("Durée totale", _minutes_to_str(duree_min)),
        ("Jalons complétés", f"{nb_done}/{len(jalons_data)}"),
    ]
    for k, v in rows_data:
        row = tbl.add_row()
        row.cells[0].text = k; row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = v

    doc.add_page_break()

    # ── 1. Description de l'incident ─────────────────────
    _heading(doc, "1. Description de l'incident", 1)
    _heading(doc, "1.1 Fait déclaré", 2)
    doc.add_paragraph(inc.fait)
    if inc.analyse:
        _heading(doc, "1.2 Analyse d'impact", 2)
        doc.add_paragraph(inc.analyse)
    if inc.moyens_engages:
        _heading(doc, "1.3 Moyens engagés", 2)
        doc.add_paragraph(inc.moyens_engages)
    if inc.actions_remediation:
        _heading(doc, "1.4 Actions de remédiation", 2)
        doc.add_paragraph(inc.actions_remediation)
    if inc.intervenant_nom:
        _heading(doc, "1.5 Intervenants", 2)
        doc.add_paragraph(f"{inc.intervenant_nom}" + (f" — {inc.intervenant_contact}" if inc.intervenant_contact else ""))

    # ── 2. Chronologie des jalons ─────────────────────────
    if jalons_data:
        _heading(doc, "2. Chronologie des jalons de résolution", 1)
        tbl2 = doc.add_table(rows=1, cols=3)
        tbl2.style = "Table Grid"
        hdr = tbl2.rows[0].cells
        hdr[0].text = "Jalon"; hdr[1].text = "Statut"; hdr[2].text = "Complété le"
        for j in jalons_data:
            r = tbl2.add_row()
            r.cells[0].text = j.get("label", "")
            r.cells[1].text = "✓ FAIT" if j.get("done") else "— En attente"
            r.cells[2].text = j.get("done_at", "")[:16] if j.get("done_at") else "—"

    # ── 3. Décisions de cellule ───────────────────────────
    if decisions:
        _heading(doc, "3. Décisions de la cellule de crise", 1)
        tbl3 = doc.add_table(rows=1, cols=4)
        tbl3.style = "Table Grid"
        h = tbl3.rows[0].cells
        h[0].text="Horodatage"; h[1].text="Responsable"; h[2].text="Décision"; h[3].text="Base régl."
        for d in decisions:
            r = tbl3.add_row()
            r.cells[0].text = d.timestamp.strftime("%d/%m %H:%M") if d.timestamp else "—"
            r.cells[1].text = d.responsable or "—"
            r.cells[2].text = d.contenu
            r.cells[3].text = d.base_reglementaire or "—"

    # ── 4. Présences cellule ──────────────────────────────
    if presences:
        _heading(doc, "4. Registre des présences", 1)
        tbl4 = doc.add_table(rows=1, cols=4)
        tbl4.style = "Table Grid"
        h = tbl4.rows[0].cells
        h[0].text="Horodatage"; h[1].text="Nom"; h[2].text="Rôle"; h[3].text="Action"
        for p in presences:
            r = tbl4.add_row()
            r.cells[0].text = p.timestamp.strftime("%d/%m %H:%M") if p.timestamp else "—"
            r.cells[1].text = p.nom
            r.cells[2].text = p.role or "—"
            r.cells[3].text = p.action

    # ── 5. Consignes de relève ────────────────────────────
    if consignes:
        _heading(doc, "5. Consignes de relève", 1)
        for c in consignes:
            p_c = doc.add_paragraph()
            p_c.add_run(f"[{c.timestamp.strftime('%d/%m %H:%M') if c.timestamp else '—'}] Pour {c.pour} : ").bold = True
            p_c.add_run(c.texte)
            if c.accuse:
                p_c.add_run(f" ✓ Accusé le {c.accuse_at.strftime('%d/%m %H:%M') if c.accuse_at else '?'}").italic = True

    # ── 6. Avis Albert ───────────────────────────────────
    if inc.albert_avis:
        _heading(doc, "6. Avis Albert AI", 1)
        doc.add_paragraph(inc.albert_avis)
        doc.add_paragraph("Note : analyse générée automatiquement par Albert AI (DINUM/Etalab)."
                          "\nDoit être interprétée par un expert métier avant toute décision.", style="Caption")

    # ── 7. Métriques ─────────────────────────────────────
    _heading(doc, "7. Métriques de gestion", 1)
    tbl_m = doc.add_table(rows=0, cols=2)
    tbl_m.style = "Table Grid"
    metrics = [
        ("Durée totale (déclaration → résolution)", _minutes_to_str(duree_min)),
        ("Jalons complétés", f"{nb_done}/{len(jalons_data)} ({int(nb_done/len(jalons_data)*100) if jalons_data else 0}%)"),
        ("Nombre de décisions actées", str(len(decisions))),
        ("Intervenants cellule", str(len(set(p.nom for p in presences)))),
        ("Consignes de relève", str(len(consignes))),
    ]
    for k, v in metrics:
        r = tbl_m.add_row()
        r.cells[0].text = k; r.cells[0].paragraphs[0].runs[0].bold = True
        r.cells[1].text = v

    # ── Pied de page ─────────────────────────────────────
    doc.add_paragraph()
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_foot.add_run(f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} — SCRIBE v5 Établissement — CONFIDENTIEL")
    run.font.size = Pt(8); run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── REX ──────────────────────────────────────────────────

class RexCreate(BaseModel):
    incident_id:     Optional[int] = None
    titre:           str
    type_crise:      Optional[str] = None
    duree_minutes:   Optional[int] = None
    nb_poles:        int = 0
    nb_decisions:    int = 0
    nb_jalons_total: int = 0
    nb_jalons_done:  int = 0
    mttd_minutes:    Optional[int] = None
    mttr_minutes:    Optional[int] = None
    points_positifs: Optional[List[str]] = []
    points_amelio:   Optional[List[str]] = []
    actions_futures: Optional[List[str]] = []
    lecons:          Optional[str] = ""
    redacteur:       Optional[str] = ""

class RexOut(BaseModel):
    id: int; incident_id: Optional[int]; titre: str; type_crise: Optional[str]
    created_at: Optional[datetime]; duree_minutes: Optional[int]
    nb_poles: int; nb_decisions: int; nb_jalons_total: int; nb_jalons_done: int
    mttd_minutes: Optional[int]; mttr_minutes: Optional[int]
    points_positifs: Optional[str]; points_amelio: Optional[str]
    actions_futures: Optional[str]; lecons: Optional[str]; redacteur: Optional[str]
    class Config: from_attributes = True


@router.get("/rapport/{incident_id}")
def download_rapport(incident_id: int, db: Session = Depends(get_db)):
    docx_bytes = generate_rapport_docx(incident_id, db)
    filename = f"rapport_cloture_incident_{incident_id}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@router.get("/rex", response_model=List[RexOut])
def list_rex(db: Session = Depends(get_db)):
    return db.query(RexEntry).order_by(RexEntry.created_at.desc()).all()

@router.post("/rex", response_model=RexOut)
def create_rex(body: RexCreate, db: Session = Depends(get_db)):
    data = body.dict()
    for k in ["points_positifs", "points_amelio", "actions_futures"]:
        if isinstance(data[k], list):
            data[k] = json.dumps(data[k], ensure_ascii=False)
    r = RexEntry(**data)
    db.add(r); db.commit(); db.refresh(r)
    return r

@router.delete("/rex/{rid}")
def delete_rex(rid: int, db: Session = Depends(get_db)):
    r = db.query(RexEntry).filter(RexEntry.id == rid).first()
    if not r: raise HTTPException(404, "REX non trouvé")
    db.delete(r); db.commit()
    return {"status": "deleted"}

@router.get("/rex-stats")
def rex_stats(db: Session = Depends(get_db)):
    entries = db.query(RexEntry).all()
    if not entries:
        return {"total": 0}
    mttrs = [e.mttr_minutes for e in entries if e.mttr_minutes]
    mttds = [e.mttd_minutes for e in entries if e.mttd_minutes]
    by_type = {}
    for e in entries:
        t = e.type_crise or "INCONNU"
        by_type[t] = by_type.get(t, 0) + 1
    return {
        "total": len(entries),
        "avg_mttr_min": int(sum(mttrs)/len(mttrs)) if mttrs else None,
        "avg_mttd_min": int(sum(mttds)/len(mttds)) if mttds else None,
        "by_type": by_type,
        "avg_jalons_pct": int(sum(e.nb_jalons_done/(e.nb_jalons_total or 1)*100 for e in entries)/len(entries)),
    }
